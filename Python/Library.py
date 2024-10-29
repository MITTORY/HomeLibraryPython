import sqlite3
import os
import csv
from docx import Document
import pandas as pd
        
class Library:
    def __init__(self):
        self.conn = sqlite3.connect('library.db')
        self.cursor = self.conn.cursor()

    def add_book(self, title, author, genre, year, book_type, description):
        self.cursor.execute(''' 
            INSERT INTO books (title, author, genre, year, type, description)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (title, author, genre, year, book_type, description))
        self.conn.commit()
        print(f'\n✅ {book_type.capitalize()} "{title}" успешно добавлена!')

    def remove_book(self, book_id):
        self.cursor.execute('SELECT * FROM books WHERE id = ?', (book_id,))
        book = self.cursor.fetchone()
        if book:
            self.cursor.execute('DELETE FROM books WHERE id = ?', (book_id,))
            self.conn.commit()
            print(f'\n✅ Книга/комикс/манга с ID {book_id} успешно удалена!')
        else:
            print(f'\n❌ Книга/комикс/манга с ID {book_id} не найдена.')

    def edit_book(self, book_id, field, new_value):
    # Проверка, является ли поле допустимым
        if field not in ["title", "author", "genre", "year", "type", "description"]:
            print("🚫 Неверное поле для редактирования.")
            return
    # Обновляем запись
        self.cursor.execute(f'UPDATE books SET {field} = ? WHERE id = ?', (new_value, book_id))
        self.conn.commit()
        print(f'\n✅ Книга/комикс/манга с ID {book_id} успешно обновлена!')

    def display_books(self):
        self.cursor.execute('SELECT * FROM books')
        books = self.cursor.fetchall()
        book_count = len(books)
        if not books:
            print("\n📚 Ваша библиотека пуста.")
            return
        print(f"\n📚 Ваша домашняя библиотека ({book_count}): ")
        print("-" * 80)
        for book in books:
            print(f'   📖 "{book[1]}"')
            print(f'   Автор: {book[2]} | Жанр: {book[3]} | Год: {book[4]} | Тип: {book[5]}')
            print(f'   Описание: {book[6] if book[6] else "Нет описания"}')
            print("-" * 80)
    
    def display_books_without_description(self):
        self.cursor.execute('SELECT * FROM books')
        books = self.cursor.fetchall()
        book_count = len(books)
        if not books:
            print("\n📚 Ваша библиотека пуста.")
            return
        print(f"\n📚 Ваша домашняя библиотека ({book_count}): ")
        print("-" * 80)
        for book in books:
            print(f'   📖 "{book[1]}"')
            print(f'   Автор: {book[2]} | Жанр: {book[3]} | Год: {book[4]} | Тип: {book[5]}')
            print("-" * 80)


    def save_to_txt(self):
        with open('library.txt', 'w', encoding='utf-8') as f:
            f.write("📚 Ваша домашняя библиотека:\n")
            f.write("-" * 80 + "\n")
            self.cursor.execute('SELECT * FROM books')
            books = self.cursor.fetchall()
            for i, book in enumerate(books, start=1):
                f.write(f'{i}. 📖 "{book[1]}"\n')
                f.write(f'   Автор: {book[2]} | Жанр: {book[3]} | Год: {book[4]} | Тип: {book[5]}\n')
                f.write(f'   Описание: {book[6] if book[6] else "Нет описания"}\n')
                f.write("-" * 80 + "\n")
        print("✅📄 Библиотека сохранена в файл library.txt.")

    def save_to_word(self):
        document = Document()
        document.add_heading('📚 Ваша домашняя библиотека:', level=1)
        self.cursor.execute('SELECT * FROM books')
        books = self.cursor.fetchall()
        for i, book in enumerate(books, start=1):
            document.add_paragraph(f'{i}. 📖 "{book[1]}"')
            document.add_paragraph(f'   Автор: {book[2]} | Жанр: {book[3]} | Год: {book[4]} | Тип: {book[5]}')
            document.add_paragraph(f'   Описание: {book[6] if book[6] else "Нет описания"}')
            document.add_paragraph("-" * 80)
        document.save('library.docx')
        print("✅📄 Библиотека сохранена в файл library.docx.")

    def save_to_excel(self):
        books = self.cursor.execute('SELECT * FROM books').fetchall()
        df = pd.DataFrame(books, columns=['ID', 'Название', 'Автор', 'Жанр', 'Год', 'Тип', 'Описание'])
        df.to_excel('library.xlsx', index=False)
        print("✅📄 Библиотека сохранена в файл library.xlsx.")

    def save_to_html(self):
        self.cursor.execute('SELECT * FROM books')
        books = self.cursor.fetchall()
        with open('library.html', 'w', encoding='utf-8') as f:
            f.write('<html><head><title>Ваша библиотека</title></head><body>')
            f.write('<h1>📚 Ваша домашняя библиотека:</h1>')
            f.write('<table border="1"><tr><th>№</th><th>Название</th><th>Автор</th><th>Жанр</th><th>Год</th><th>Тип</th><th>Описание</th></tr>')
            for i, book in enumerate(books, start=1):
                f.write(f'<tr><td>{i}</td><td>{book[1]}</td><td>{book[2]}</td><td>{book[3]}</td><td>{book[4]}</td><td>{book[5]}</td><td>{book[6] if book[6] else "Нет описания"}</td></tr>')
            f.write('</table></body></html>')
        print("✅📄 Библиотека сохранена в файл library.html.")

    def search_books(self, search_type, search_value):
        if search_type == "title":
            self.cursor.execute('SELECT * FROM books WHERE title LIKE ? ORDER BY title ASC', (f'%{search_value}%',))
        elif search_type == "author":
            self.cursor.execute('SELECT * FROM books WHERE author LIKE ? ORDER BY author ASC', (f'%{search_value}%',))
        elif search_type == "year":
            # Чтобы сортировать книги по году, используем условие поиска и сортировку
            self.cursor.execute('SELECT * FROM books WHERE year = ? ORDER BY year ASC', (search_value,))
        elif search_type == "type":
            self.cursor.execute('SELECT * FROM books WHERE type LIKE ? ORDER BY type ASC', (f'%{search_value}%',))
        elif search_type == "genre":
            self.cursor.execute('SELECT * FROM books WHERE genre LIKE ? ORDER BY genre ASC', (f'%{search_value}%',))
        else:
            print("🚫 Неверный тип поиска.")
            return
    
        books = self.cursor.fetchall()
        if not books:
            print("\n❌ Книги не найдены.")
            return
    
        print("\n📚 Результаты поиска:")
        print("-" * 60)
        for i, book in enumerate(books, start=1):
            print(f'{i}. 📖 "{book[1]}"')
            print(f'   Автор: {book[2]} | Жанр: {book[3]} | Год: {book[4]} | Тип: {book[5]}')
            print("-" * 60)


    def get_unique_values(self, column):
        self.cursor.execute(f'SELECT DISTINCT {column} FROM books')
        return [row[0] for row in self.cursor.fetchall()]

    def display_search_options(self, column):
        self.cursor.execute(f'SELECT DISTINCT {column} FROM books ORDER BY {column} ASC')
        values = [row[0] for row in self.cursor.fetchall()]
        print(f"\nДоступно:")
        for value in values:
            print(f"- {value}")


    def import_books(self, filename):
        with open(filename, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            next(reader)
            for row in reader:
                self.add_book(row[0], row[1], row[2], int(row[3]), row[4])

    def export_books(self):
        with open('exported_books.csv', 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Название', 'Автор', 'Жанр', 'Год', 'Тип'])
            self.cursor.execute('SELECT * FROM books')
            books = self.cursor.fetchall()
            for book in books:
                writer.writerow([book[1], book[2], book[3], book[4], book[5]])
        print("📤 Библиотека экспортирована в файл exported_books.csv.")

    def menu(self):
        while True:
            print("\n" + "=" * 40)
            print("🏠 Добро пожаловать в библиотеку! 📚")
            print("=" * 40)
            print("1. 🔍 Поиск")
            print("2. 📜 Показать библиотеку (с описанием)")
            print("3. 📜 Показать библиотеку (без описания)")
            print("4. ➕ Добавить книгу/комикс/мангу")
            print("5. ➖ Удалить книгу/комикс/мангу")
            print("6. ✏️ Редактировать книгу/комикс/мангу")
            print("7. 📥 Импорт библиотеки")
            print("8. 📤 Экспорт библиотеки")
            print("9. 💾 Сохранить библиотеку")
            print("0. 🚪 Покинуть библиотеку")
            print("=" * 40)

            choice = input("👉 Выберите действие (1-9): ")
            if choice == "1":
                print("\n🔍 Поиск по:")
                print("1. Названию")
                print("2. Автору")
                print("3. Жанру")
                print("4. Году")
                print("5. Типу")
                search_choice = input("👉 Выберите вариант (1-5): ")
                if search_choice == "1":
                    search_value = input("Введите название: ")
                    self.search_books("title", search_value)
                elif search_choice == "2":
                    self.display_search_options("author")
                    search_value = input("Введите автора: ")
                    self.search_books("author", search_value)
                elif search_choice == "3":
                    self.display_search_options("genre")
                    search_value = input("Введите жанр: ")
                    self.search_books("genre", search_value)
                elif search_choice == "4":
                    self.display_search_options("year")
                    search_value = input("Введите год: ")
                    self.search_books("year", search_value)
                elif search_choice == "5":
                    self.display_search_options("type")
                    search_value = input("Введите тип: ")
                    self.search_books("type", search_value)
                else:
                    print("🚫 Неверный вариант.")

            elif choice == "2":
                self.display_books()
                
            elif choice == "3":
                self.display_books_without_description()

            elif choice == "4":
                title = input("Введите название: ")
                author = input("Введите автора: ")
                genre = input("Введите жанр: ")
                year = int(input("Введите год: "))
                book_type = input("Введите тип (книга/комикс/манга): ")
                description = input("Введите описание (необязательно): ")
                self.add_book(title, author, genre, year, book_type, description)

            elif choice == "5":
                try:
                    book_id = int(input("Введите ID книги/комикса/манги для удаления: "))
                    self.remove_book(book_id)
                except ValueError:
                    print("🚫 Неверный формат ID. Пожалуйста, введите числовое значение.")

            elif choice == "6":
                book_id = int(input("Введите ID книги/комикса/манги для редактирования: "))
                print("\nВыберите поле для редактирования:")
                print("1. Название")
                print("2. Автор")
                print("3. Жанр")
                print("4. Год")
                print("5. Тип")
                print("6. Описание")
                field_choice = input("👉 Выберите поле (1-6): ")
                fields = ["title", "author", "genre", "year", "type", "description"]
            
                if field_choice in [str(i + 1) for i in range(len(fields))]:
                    field = fields[int(field_choice) - 1]
                    new_value = input(f"Введите новое значение для {field}: ")
                    if field == "year":
                        new_value = int(new_value)
                    self.edit_book(book_id, field, new_value)
                else:
                    print("🚫 Неверный вариант.")

            elif choice == "7":
                filename = input("Введите имя файла для импорта (например, books.csv): ")
                self.import_books(filename)

            elif choice == "8":
                self.export_books()

            elif choice == "9":
                print("\nВыберите формат для сохранения:")
                print("1. Сохранить в TXT")
                print("2. Сохранить в Word")
                print("3. Сохранить в Excel")
                print("4. Сохранить в HTML")
                print("5. Сохранить все форматы")
                format_choice = input("👉 Выберите вариант (1-5): ")
    
                if format_choice == "1":
                    self.save_to_txt()
                elif format_choice == "2":
                    self.save_to_word()
                elif format_choice == "3":
                    self.save_to_excel()
                elif format_choice == "4":
                    self.save_to_html()
                elif format_choice == "5":
                    self.save_to_txt()
                    self.save_to_word()
                    self.save_to_excel()
                    self.save_to_html()
                else:
                    print("🚫 Неверный вариант.")
                    
            elif choice == "0":
                print("👋 Спасибо за использование нашей библиотеки! До свидания!")
                break

            else:
                print("🚫 Неверный вариант.")
    
if __name__ == "__main__":
    library = Library()
    library.menu()
